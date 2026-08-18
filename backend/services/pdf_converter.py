"""
Modern High-Performance PDF to Word Converter Service
Optimized specifically for Vietnamese Math Exams (Đề toán 12),
Complex Formulas, Geometry Diagrams, Tables, and General Documents.
Includes Smart Vector & Raster Diagram Cropper (Bảng biến thiên, Đồ thị hàm số, Hình học không gian).
Supports live text extraction, editing, and saving back to .docx with embedded diagrams.
"""

import os
import re
import glob
import time
import tempfile
import logging
import asyncio
from pathlib import Path
from concurrent.futures import ThreadPoolExecutor
import pymupdf  # Modern PyMuPDF (fitz)
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION

logger = logging.getLogger(__name__)

# Dedicated thread pool for heavy PDF processing (prevents blocking FastAPI event loop)
executor = ThreadPoolExecutor(max_workers=4)


def detect_and_crop_diagrams(page, page_idx: int = 0, dpi: int = 300, output_dir: str = None) -> list:
    """
    SMART DIAGRAM, GRAPH & VARIATION TABLE DETECTOR:
    Identifies vector drawing clusters (Bảng biến thiên, Đồ thị hàm số, Hình học không gian)
    and raster images, merges nearby shapes, and crops them as ultra-crisp 300 DPI PNGs.
    """
    pw, ph = page.rect.width, page.rect.height
    drawings = page.get_drawings()
    raw_rects = []

    # 1. Collect valid vector drawing rectangles
    for d in drawings:
        r = d["rect"]
        # Skip full-page border lines or full-width separators
        if r.width > 0.92 * pw and r.height < 3:
            continue
        if r.width > 0.92 * pw and r.height > 0.92 * ph:
            continue
        if r.width < 4 and r.height < 4:
            continue
        raw_rects.append(r)

    # 2. Collect raster image rectangles
    images = page.get_images()
    for img_info in images:
        xref = img_info[0]
        for ir in page.get_image_rects(xref):
            if ir.width < 0.92 * pw and ir.height < 0.92 * ph and (ir.width * ir.height) >= 400:
                raw_rects.append(ir)

    if not raw_rects:
        return []

    # 3. Cluster overlapping & closely situated drawing elements
    thresh_x = 24
    thresh_y = 20
    clusters = []

    for r in raw_rects:
        matched = []
        cur_box = pymupdf.Rect(r)
        for i, c in enumerate(clusters):
            expanded_c = pymupdf.Rect(c.x0 - thresh_x, c.y0 - thresh_y, c.x1 + thresh_x, c.y1 + thresh_y)
            if expanded_c.intersects(cur_box):
                matched.append(i)

        if not matched:
            clusters.append(cur_box)
        else:
            new_box = cur_box
            for idx in matched:
                new_box = new_box | clusters[idx]
            clusters = [c for i, c in enumerate(clusters) if i not in matched]
            clusters.append(new_box)

    # 4. Filter and crop meaningful diagrams (graphs, variation tables, geometry diagrams)
    if not output_dir:
        output_dir = tempfile.gettempdir()

    diagrams = []
    diag_idx = 1

    # Sort clusters from top to bottom
    clusters.sort(key=lambda c: (c.y0, c.x0))

    for c in clusters:
        # Minimum size threshold for a diagram/graph/variation table
        if c.width >= 35 and c.height >= 25 and (c.width * c.height) >= 900:
            # Add safe padding to avoid clipping labels/arrows
            pad_x = 8
            pad_y = 8
            padded = pymupdf.Rect(
                max(0, c.x0 - pad_x),
                max(0, c.y0 - pad_y),
                min(pw, c.x1 + pad_x),
                min(ph, c.y1 + pad_y)
            )

            # High resolution rendering crop (300 DPI)
            zoom = dpi / 72.0
            mat = pymupdf.Matrix(zoom, zoom)
            pix = page.get_pixmap(matrix=mat, clip=padded, alpha=False)

            temp_img_path = os.path.join(output_dir, f"diag_p{page_idx + 1}_{diag_idx}_{os.getpid()}.png")
            pix.save(temp_img_path)

            width_in = min(5.5, max(2.5, round(padded.width / 72.0, 2)))
            tag = f"[🖼️ Đồ thị / Hình vẽ trang {page_idx + 1} - Hình {diag_idx}]"

            diagrams.append({
                "rect": padded,
                "x0": padded.x0,
                "y0": padded.y0,
                "x1": padded.x1,
                "y1": padded.y1,
                "img_path": temp_img_path,
                "width_in": width_in,
                "tag": tag,
                "page_idx": page_idx,
                "diag_idx": diag_idx,
            })
            diag_idx += 1

    return diagrams


def extract_text_from_pdf(pdf_path: str, pages_to_convert: list[int] = None) -> str:
    """Extract clean, structured text with diagram markers from PDF pages."""
    try:
        doc = pymupdf.open(pdf_path)
        total_pages = len(doc)
        if not pages_to_convert:
            pages = list(range(total_pages))
        else:
            pages = [p for p in pages_to_convert if 0 <= p < total_pages]

        page_texts = []
        for p_idx in pages:
            page = doc[p_idx]
            diagrams = detect_and_crop_diagrams(page, p_idx, dpi=200)
            blocks = page.get_text("blocks")

            items = []
            for b in blocks:
                text = b[4].strip()
                if not text:
                    continue
                # If text is small and completely inside a diagram (like axis numbers), don't clutter text
                b_rect = pymupdf.Rect(b[:4])
                inside = False
                for d in diagrams:
                    if d['rect'].contains(b_rect) and len(text) <= 5:
                        inside = True
                        break
                if not inside:
                    items.append({"y0": b[1], "content": text, "type": "text"})

            for d in diagrams:
                items.append({"y0": d["y0"], "content": d["tag"], "type": "diagram"})

            items.sort(key=lambda x: x["y0"])

            page_content = [it["content"] for it in items]
            if page_content:
                page_texts.append(f"--- [ Trang {p_idx + 1} ] ---\n" + "\n\n".join(page_content))
            else:
                page_texts.append(f"--- [ Trang {p_idx + 1} ] ---\n(Trang này chủ yếu là hình ảnh hoặc sơ đồ)")

        doc.close()
        return "\n\n".join(page_texts)
    except Exception as e:
        logger.warning(f"Error extracting text from PDF: {e}")
        return ""


def extract_text_from_docx(docx_path: str) -> str:
    """Extract text paragraphs from a Word (.docx) file."""
    try:
        doc = Document(docx_path)
        texts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        if texts:
            return "\n\n".join(texts)
        return ""
    except Exception as e:
        logger.warning(f"Error extracting text from docx: {e}")
        return ""


def find_diagram_image_for_tag(tag: str) -> str:
    """Find the cached cropped diagram image file on disk matching a tag."""
    # Match patterns like: [🖼️ Đồ thị / Hình vẽ trang 1 - Hình 2]
    m = re.search(r"trang\s*(\d+).*?Hình\s*(\d+)", tag, re.IGNORECASE)
    if m:
        p_num, img_num = m.group(1), m.group(2)
        pattern = os.path.join(tempfile.gettempdir(), f"diag_p{p_num}_{img_num}_*.png")
        matches = glob.glob(pattern)
        if matches:
            return matches[-1]

    # Direct file path tag support
    m_path = re.search(r"\[IMAGE_PATH:\s*(.*?)\]", tag)
    if m_path:
        p = m_path.group(1).strip()
        if os.path.exists(p):
            return p

    return None


def save_text_to_docx(docx_path: str, text: str) -> dict:
    """
    Format and save user-edited text into a clean, beautifully styled Word (.docx) document,
    re-embedding any cropped diagram images (bảng biến thiên, đồ thị) positioned at [🖼️ ...].
    """
    word_doc = Document()

    # Standard A4 Margins: Top/Bottom/Left/Right = 0.75 in (54 pt)
    margin_pt = Pt(54)
    for section in word_doc.sections:
        section.top_margin = margin_pt
        section.bottom_margin = margin_pt
        section.left_margin = margin_pt
        section.right_margin = margin_pt

    lines = text.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        if not line:
            i += 1
            continue

        # Page separator marker
        if re.match(r"^---\s*\[\s*Trang\s*\d+\s*\]\s*---", line, re.IGNORECASE):
            p = word_doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run(line)
            run.font.name = "Times New Roman"
            run.font.size = Pt(10)
            run.font.italic = True
            run.font.color.rgb = RGBColor(140, 120, 90)
            i += 1
            continue

        # Diagram / Graph Image Placeholder Tag: [🖼️ ...]
        if line.startswith("[🖼️") or "[IMAGE_PATH:" in line:
            img_file = find_diagram_image_for_tag(line)
            if img_file and os.path.exists(img_file):
                p = word_doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(8)
                run = p.add_run()
                run.add_picture(img_file, width=Inches(4.5))
            else:
                # Fallback text note if image not found
                p = word_doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(line)
                run.font.name = "Times New Roman"
                run.font.size = Pt(9.5)
                run.font.italic = True
                run.font.color.rgb = RGBColor(100, 100, 100)
            i += 1
            continue

        # Big Headings: # Heading or PHẦN / CHƯƠNG / BÀI / ĐỀ THI
        if line.startswith("# ") or re.match(r"^(ĐỀ THI|BÀI THI|CHƯƠNG|PHẦN|BỘ GIÁO DỤC|TRƯỜNG)\b", line, re.IGNORECASE):
            heading_text = line.lstrip("# ").strip()
            p = word_doc.add_paragraph()
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(4)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if ("ĐỀ" in line.upper() or "TRƯỜNG" in line.upper()) else WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(heading_text)
            run.font.name = "Times New Roman"
            run.font.size = Pt(14)
            run.bold = True
            run.font.color.rgb = RGBColor(30, 41, 59)
            i += 1
            continue

        # Sub-headings: ## or ### or Câu / Question
        if line.startswith("## ") or line.startswith("### ") or re.match(r"^(Câu|Bài|Question|Part)\s*\d+[\.\:]?", line, re.IGNORECASE):
            heading_text = line.lstrip("# ").strip()
            p = word_doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(3)
            run = p.add_run(heading_text)
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)
            run.bold = True
            run.font.color.rgb = RGBColor(15, 23, 42)
            i += 1
            continue

        # Bullet list items
        if line.startswith("- ") or line.startswith("* ") or line.startswith("• "):
            p = word_doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.line_spacing = 1.15
            run = p.add_run(line[2:].strip())
            run.font.name = "Times New Roman"
            run.font.size = Pt(11)
            i += 1
            continue

        # Standard Paragraph
        p = word_doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.2
        run = p.add_run(line)
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(20, 20, 20)
        i += 1

    word_doc.save(docx_path)
    file_size = os.path.getsize(docx_path)
    words = len(text.split())
    chars = len(text)

    return {
        "success": True,
        "output_size": file_size,
        "word_count": words,
        "char_count": chars,
        "message": f"Đã lưu thành công {words} từ và hình ảnh vào file Word"
    }


def get_pdf_info(pdf_path: str) -> dict:
    """
    Extract PDF metadata, page count, and detect whether it's math/scanned/text.
    """
    doc = pymupdf.open(pdf_path)
    total_pages = len(doc)
    
    total_words = 0
    total_images = 0
    total_drawings = 0
    sample_pages = []

    for i in range(min(total_pages, 5)):
        page = doc[i]
        text = page.get_text("text").strip()
        word_count = len(text.split())
        images = page.get_images()
        drawings = page.get_drawings()
        
        total_words += word_count
        total_images += len(images)
        total_drawings += len(drawings)

        sample_pages.append({
            "page_num": i + 1,
            "width": round(page.rect.width, 1),
            "height": round(page.rect.height, 1),
            "word_count": word_count,
            "image_count": len(images),
            "drawing_count": len(drawings),
            "text_preview": text[:250].strip() if text else "(Không có text / Chứa công thức vector hoặc ảnh scan)",
        })

    # Detect PDF nature
    is_scanned = (total_words < 10 and total_images > 0)
    is_math_heavy = (total_drawings > 20 or total_images > 3 or ("toán" in str(doc.metadata).lower()))

    doc.close()
    
    return {
        "num_pages": total_pages,
        "is_scanned": is_scanned,
        "is_math_heavy": is_math_heavy,
        "recommended_mode": "hybrid" if is_math_heavy else "editable_text",
        "metadata": doc.metadata if hasattr(doc, 'metadata') else {},
        "pages": sample_pages,
    }


def generate_pdf_preview(pdf_path: str, page_num: int = 0, dpi: int = 150) -> bytes:
    """Generate high quality PNG preview of any PDF page."""
    doc = pymupdf.open(pdf_path)
    if page_num < 0 or page_num >= len(doc):
        doc.close()
        raise ValueError(f"Trang {page_num + 1} không tồn tại (tổng {len(doc)} trang)")

    page = doc[page_num]
    zoom = dpi / 72.0
    mat = pymupdf.Matrix(zoom, zoom)
    pix = page.get_pixmap(matrix=mat, alpha=False)
    img_bytes = pix.tobytes("png")
    doc.close()
    return img_bytes


def convert_math_hd_mode(pdf_path: str, output_path: str, pages_to_convert: list[int] = None, dpi: int = 250) -> dict:
    """
    ULTRA-HD MATH & EXAM ENGINE:
    Renders pages at ultra-high DPI with exact vector & formula fidelity.
    Preserves 100% integrals, fractions, geometry diagrams, coordinate axes,
    and table variations without any corrupted characters or missing formulas.
    Also extracts text with diagram markers for viewing & editing in the frontend editor.
    """
    start_time = time.time()
    doc = pymupdf.open(pdf_path)
    total_pages = len(doc)

    if not pages_to_convert:
        pages_to_convert = list(range(total_pages))
    else:
        pages_to_convert = [p for p in pages_to_convert if 0 <= p < total_pages]

    word_doc = Document()

    # Configure margins to 0.4 inch (minimal margins for maximum exam content area)
    margin_pt = Pt(28)
    for section in word_doc.sections:
        section.top_margin = margin_pt
        section.bottom_margin = margin_pt
        section.left_margin = margin_pt
        section.right_margin = margin_pt

    temp_images = []
    
    try:
        for idx, page_idx in enumerate(pages_to_convert):
            page = doc[page_idx]

            # High resolution rendering (250 DPI = sharp crisp text and formulas)
            zoom = dpi / 72.0
            mat = pymupdf.Matrix(zoom, zoom)
            pix = page.get_pixmap(matrix=mat, alpha=False)

            # Temp file for Word insertion
            temp_img = os.path.join(tempfile.gettempdir(), f"exam_page_{page_idx}_{os.getpid()}_{idx}.png")
            pix.save(temp_img)
            temp_images.append(temp_img)

            # Calculate proper page width & height in inches for A4 Word doc
            page_w_in = page.rect.width / 72.0
            page_h_in = page.rect.height / 72.0

            max_w_in = 7.4
            max_h_in = 10.5

            scale = min(max_w_in / page_w_in, max_h_in / page_h_in, 1.0)
            img_width = page_w_in * scale

            # Add page break before (except first page)
            if idx > 0:
                word_doc.add_page_break()

            # Insert picture centered
            p = word_doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            run = p.add_run()
            run.add_picture(temp_img, width=Inches(img_width))

        doc.close()
        word_doc.save(output_path)
        
        elapsed = round(time.time() - start_time, 2)
        file_size = os.path.getsize(output_path)

        # Extract text for editor preview
        extracted_text = extract_text_from_pdf(pdf_path, pages_to_convert)

        return {
            "success": True,
            "pages_converted": len(pages_to_convert),
            "total_pages": total_pages,
            "output_size": file_size,
            "method": "math_hd",
            "dpi": dpi,
            "elapsed_seconds": elapsed,
            "extracted_text": extracted_text,
            "word_count": len(extracted_text.split()),
            "char_count": len(extracted_text),
            "message": f"Chuyển đổi thành công {len(pages_to_convert)} trang đề toán siêu nét trong {elapsed}s"
        }
    finally:
        # Clean up temporary images
        for img_path in temp_images:
            try:
                if os.path.exists(img_path):
                    os.remove(img_path)
            except Exception:
                pass


def convert_hybrid_mode(pdf_path: str, output_path: str, pages_to_convert: list[int] = None, dpi: int = 300) -> dict:
    """
    SMART HYBRID ENGINE (VĂN BẢN + TỰ CẮT ĐỒ THỊ & BẢNG BIẾN THIÊN):
    Extracts text paragraphs as native editable Word text, automatically detects & crops
    all math graphs, variation tables (bảng biến thiên), geometry diagrams, and illustrations,
    and inserts them directly below their corresponding question in Word!
    """
    start_time = time.time()
    doc = pymupdf.open(pdf_path)
    total_pages = len(doc)

    if not pages_to_convert:
        pages_to_convert = list(range(total_pages))
    else:
        pages_to_convert = [p for p in pages_to_convert if 0 <= p < total_pages]

    word_doc = Document()
    
    # Standard margins (0.65 in)
    for section in word_doc.sections:
        section.top_margin = Pt(46)
        section.bottom_margin = Pt(46)
        section.left_margin = Pt(46)
        section.right_margin = Pt(46)

    accumulated_text = []

    try:
        for idx, page_idx in enumerate(pages_to_convert):
            page = doc[page_idx]
            
            if idx > 0:
                word_doc.add_page_break()

            # Page header
            hdr = word_doc.add_paragraph()
            hdr_run = hdr.add_run(f"— TRANG {page_idx + 1} —")
            hdr_run.font.size = Pt(9)
            hdr_run.font.color.rgb = RGBColor(140, 120, 90)
            hdr_run.italic = True
            hdr.alignment = WD_ALIGN_PARAGRAPH.CENTER

            accumulated_text.append(f"--- [ Trang {page_idx + 1} ] ---")

            # 1. Detect & Crop all graphs, variation tables and diagrams
            diagrams = detect_and_crop_diagrams(page, page_idx, dpi=dpi)

            # 2. Extract text blocks
            blocks = page.get_text("blocks")

            items = []
            for b in blocks:
                text = b[4].strip()
                if not text:
                    continue
                # If text is small and completely inside a diagram (like axis numbers), don't duplicate outside
                b_rect = pymupdf.Rect(b[:4])
                inside = False
                for d in diagrams:
                    if d['rect'].contains(b_rect) and len(text) <= 5:
                        inside = True
                        break
                if not inside:
                    items.append({"y0": b[1], "content": text, "type": "text"})

            for d in diagrams:
                items.append({"y0": d["y0"], "content": d, "type": "diagram"})

            # 3. Sort items from top to bottom
            items.sort(key=lambda x: x["y0"])

            # 4. Insert into Word in correct reading order
            for it in items:
                if it["type"] == "text":
                    text_str = it["content"]
                    accumulated_text.append(text_str)

                    # Determine styling
                    if re.match(r"^(ĐỀ THI|BÀI THI|CHƯƠNG|PHẦN|BỘ GIÁO DỤC|TRƯỜNG)\b", text_str, re.IGNORECASE):
                        p = word_doc.add_paragraph()
                        p.paragraph_format.space_before = Pt(8)
                        p.paragraph_format.space_after = Pt(4)
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if ("ĐỀ" in text_str.upper() or "TRƯỜNG" in text_str.upper()) else WD_ALIGN_PARAGRAPH.LEFT
                        run = p.add_run(text_str)
                        run.font.name = "Times New Roman"
                        run.font.size = Pt(13)
                        run.bold = True
                    elif re.match(r"^(Câu|Bài|Question|Part)\s*\d+[\.\:]?", text_str, re.IGNORECASE):
                        p = word_doc.add_paragraph()
                        p.paragraph_format.space_before = Pt(6)
                        p.paragraph_format.space_after = Pt(3)
                        run = p.add_run(text_str)
                        run.font.name = "Times New Roman"
                        run.font.size = Pt(11.5)
                        run.bold = True
                    else:
                        p = word_doc.add_paragraph()
                        p.paragraph_format.space_after = Pt(3)
                        p.paragraph_format.line_spacing = 1.15
                        run = p.add_run(text_str)
                        run.font.name = "Times New Roman"
                        run.font.size = Pt(11)

                elif it["type"] == "diagram":
                    d = it["content"]
                    accumulated_text.append(d["tag"])

                    # Insert cropped diagram picture centered in Word
                    p = word_doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p.paragraph_format.space_before = Pt(4)
                    p.paragraph_format.space_after = Pt(6)
                    run = p.add_run()
                    run.add_picture(d["img_path"], width=Inches(d["width_in"]))

        doc.close()
        word_doc.save(output_path)
        
        elapsed = round(time.time() - start_time, 2)
        file_size = os.path.getsize(output_path)
        full_text = "\n\n".join(accumulated_text)

        return {
            "success": True,
            "pages_converted": len(pages_to_convert),
            "total_pages": total_pages,
            "output_size": file_size,
            "method": "hybrid",
            "elapsed_seconds": elapsed,
            "extracted_text": full_text,
            "word_count": len(full_text.split()),
            "char_count": len(full_text),
            "message": f"Chuyển đổi thành công {len(pages_to_convert)} trang kèm tự động cắt đồ thị & bảng biến thiên ({elapsed}s)"
        }
    except Exception as e:
        logger.error(f"Error in hybrid conversion: {e}")
        # Fallback to Math HD mode if hybrid fails
        return convert_math_hd_mode(pdf_path, output_path, pages_to_convert, dpi=200)


def convert_editable_text_mode(pdf_path: str, output_path: str, pages_to_convert: list[int] = None) -> dict:
    """
    EDITABLE TEXT & TABLES ENGINE:
    Uses pdf2docx with safe parameters, falls back gracefully if pdf2docx fails.
    Best for standard documents with editable text paragraphs and tables.
    """
    start_time = time.time()
    doc = pymupdf.open(pdf_path)
    total_pages = len(doc)
    doc.close()

    try:
        from pdf2docx import Converter
        cv = Converter(pdf_path)
        
        start_page = 0
        end_page = total_pages
        if pages_to_convert:
            start_page = min(pages_to_convert)
            end_page = max(pages_to_convert) + 1
            
        cv.convert(output_path, start=start_page, end=end_page)
        cv.close()

        elapsed = round(time.time() - start_time, 2)
        file_size = os.path.getsize(output_path)

        # Extract text for preview
        extracted_text = extract_text_from_docx(output_path)
        if not extracted_text:
            extracted_text = extract_text_from_pdf(pdf_path, pages_to_convert)

        return {
            "success": True,
            "pages_converted": end_page - start_page,
            "total_pages": total_pages,
            "output_size": file_size,
            "method": "editable_text",
            "elapsed_seconds": elapsed,
            "extracted_text": extracted_text,
            "word_count": len(extracted_text.split()),
            "char_count": len(extracted_text),
            "message": f"Trích xuất văn bản chỉnh sửa được thành công ({elapsed}s)"
        }
    except Exception as e:
        logger.warning(f"pdf2docx failed or timed out: {e}. Falling back to Smart Hybrid mode...")
        return convert_hybrid_mode(pdf_path, output_path, pages_to_convert, dpi=250)


async def run_conversion_async(pdf_path: str, output_path: str, mode: str = "hybrid", pages: list[int] = None, dpi: int = 250) -> dict:
    """Run conversion in thread pool to ensure non-blocking execution."""
    loop = asyncio.get_running_loop()
    
    if mode == "hybrid":
        return await loop.run_in_executor(executor, convert_hybrid_mode, pdf_path, output_path, pages, dpi)
    elif mode == "math_hd":
        return await loop.run_in_executor(executor, convert_math_hd_mode, pdf_path, output_path, pages, dpi)
    elif mode == "editable_text":
        return await loop.run_in_executor(executor, convert_editable_text_mode, pdf_path, output_path, pages)
    else:
        return await loop.run_in_executor(executor, convert_hybrid_mode, pdf_path, output_path, pages, dpi)


async def run_save_text_async(docx_path: str, text: str) -> dict:
    """Save edited text to docx in thread pool with diagram re-embedding."""
    loop = asyncio.get_running_loop()
    return await loop.run_in_executor(executor, save_text_to_docx, docx_path, text)
